from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

# Use a system-wide Unicode font rather than a downloadable macOS font asset.
# This gives Word and headless LibreOffice a stable CJK fallback.
TASK_SHEET_CJK_FONT = "Arial Unicode MS"


def _configure(doc: Document, title: str, version: str):
    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    header = section.header.paragraphs[0]
    header.text = f"LessonForge AI · {title} · {version}"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("第 ")
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    run._r.addnext(fld)
    footer.add_run(" 页")
    styles = doc.styles
    styles["Normal"].font.name = "Microsoft YaHei"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    styles["Normal"].font.size = Pt(10.5)


def render_markdown_docx(title: str, markdown: str, output: Path, version: str = "V1") -> Path:
    doc = Document()
    _configure(doc, title, version)
    doc.add_heading(title, 0)
    doc.add_paragraph(f"版本：{version}")
    for line in markdown.splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("- "):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        elif stripped.startswith("> "):
            paragraph = doc.add_paragraph(stripped[2:])
            paragraph.style = "Quote"
        else:
            doc.add_paragraph(stripped.replace("**", ""))
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)
    Document(output)
    return output


def _exercise_answer_text(item: dict) -> str:
    answer = item.get("answer_key") or {}
    values = answer.get("correct_option_ids") or answer.get("accepted_answers") or []
    return "、".join(values) or answer.get("reference_answer", "")


def _add_exercise_stimulus(doc: Document, stimulus: dict, asset_paths: dict[str, str]) -> None:
    title = stimulus.get("title") or "材料"
    if stimulus.get("kind") == "text":
        paragraph = doc.add_paragraph()
        paragraph.style = "Quote"
        paragraph.add_run(f"{title}：").bold = True
        paragraph.add_run(stimulus.get("text", ""))
    elif stimulus.get("kind") == "table":
        doc.add_paragraph(title).runs[0].bold = True
        columns = stimulus.get("columns") or []
        rows = stimulus.get("rows") or []
        table = doc.add_table(rows=1 + len(rows), cols=len(columns))
        table.style = "Table Grid"
        for index, column in enumerate(columns):
            _shade_cell(table.cell(0, index), "E8EEF5")
            _write_cell(table.cell(0, index), str(column), bold=True, color="1F4D78")
        for row_index, values in enumerate(rows, 1):
            for column_index, value in enumerate(values):
                _write_cell(table.cell(row_index, column_index), str(value))
        if columns:
            base = 9360 // len(columns)
            widths = [base] * len(columns)
            widths[-1] += 9360 - sum(widths)
            _set_table_geometry(table, widths)
    elif stimulus.get("visual"):
        visual = stimulus["visual"]
        asset_path = asset_paths.get(visual.get("asset_id", ""))
        if asset_path and Path(asset_path).is_file():
            try:
                doc.add_picture(asset_path, width=Inches(5.8))
            except Exception:
                doc.add_paragraph(f"配图说明：{visual.get('alt_text', '')}")
        else:
            doc.add_paragraph(f"图示说明：{visual.get('fallback_stimulus') or visual.get('alt_text', '')}")
        if visual.get("caption"):
            caption = doc.add_paragraph(visual["caption"])
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _add_exercise_question(doc: Document, item: dict, number: int, include_answers: bool) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(f"{number}. {item['stem']}  ")
    run.bold = True
    paragraph.add_run(f"（{item['score']} 分）")
    for option in item.get("options", []):
        doc.add_paragraph(f"{option['id']}. {option['text']}", style="List Bullet")

    if not include_answers:
        space = item.get("answer_space") or {}
        if space.get("mode") == "table" and space.get("columns"):
            columns = space["columns"]
            table = doc.add_table(rows=1 + max(1, space.get("blank_rows", 1)), cols=len(columns))
            table.style = "Table Grid"
            for index, column in enumerate(columns):
                _shade_cell(table.cell(0, index), "F2F4F7")
                _write_cell(table.cell(0, index), column, bold=True)
            for row in table.rows[1:]:
                row.height = Inches(0.38)
                row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
            base = 9360 // len(columns)
            widths = [base] * len(columns)
            widths[-1] += 9360 - sum(widths)
            _set_table_geometry(table, widths)
        elif space.get("mode") in {"lines", "grid"}:
            for _ in range(max(1, space.get("lines", 2))):
                line = doc.add_paragraph(" ")
                line.paragraph_format.space_after = Pt(5)
                p_pr = line._p.get_or_add_pPr()
                p_bdr = p_pr.find(qn("w:pBdr"))
                if p_bdr is None:
                    p_bdr = OxmlElement("w:pBdr")
                    p_pr.append(p_bdr)
                bottom = OxmlElement("w:bottom")
                bottom.set(qn("w:val"), "single")
                bottom.set(qn("w:sz"), "4")
                bottom.set(qn("w:color"), "94A3B8")
                p_bdr.append(bottom)
        return

    def answer_paragraph():
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0.12)
        paragraph.paragraph_format.right_indent = Inches(0.08)
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        p_pr = paragraph._p.get_or_add_pPr()
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), "EEF2FF")
        p_pr.append(shading)
        borders = OxmlElement("w:pBdr")
        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single")
        left.set(qn("w:sz"), "18")
        left.set(qn("w:color"), "4F46E5")
        borders.append(left)
        p_pr.append(borders)
        return paragraph

    paragraph = answer_paragraph()
    paragraph.add_run("参考答案：").bold = True
    paragraph.add_run(_exercise_answer_text(item))
    paragraph = answer_paragraph()
    paragraph.add_run("解析：").bold = True
    paragraph.add_run(item.get("analysis", ""))
    if item.get("scoring_points"):
        paragraph = answer_paragraph()
        paragraph.add_run("评分点：").bold = True
        for point in item["scoring_points"]:
            paragraph = answer_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.35)
            paragraph.add_run(f"{point['id']} · {point['criterion']}（{point['points']} 分）：{point['acceptable_evidence']}")
    if item.get("common_errors"):
        paragraph = answer_paragraph()
        paragraph.add_run("常见错误：").bold = True
        paragraph.add_run("；".join(item["common_errors"]))


def render_exercise_docx(
    title: str,
    content: dict,
    output: Path,
    include_answers: bool,
    asset_paths: dict[str, str] | None = None,
) -> Path:
    if content.get("schema_version") != "2.0":
        doc = Document()
        _configure(doc, title, "V1")
        doc.add_heading(title, 0)
        for item in content.get("items", []):
            doc.add_heading(f"{item['id']} · {item['stem']}", level=2)
            for option in item.get("options", []):
                doc.add_paragraph(option)
            if include_answers:
                doc.add_paragraph(f"答案：{'、'.join(item.get('correct_answers', []))}")
                doc.add_paragraph(f"解析：{item.get('explanation', '')}")
        output.parent.mkdir(parents=True, exist_ok=True)
        doc.save(output)
        Document(output)
        return output

    asset_paths = asset_paths or {}
    doc = Document()
    _configure_task_sheet(doc, title, "V2", "教师版" if include_answers else "学生版")
    settings = content["paper_settings"]
    info = content["course_info"]
    doc.add_heading(settings["title"], 0)
    subtitle = doc.add_paragraph("教师版" if include_answers else "学生版", style="Subtitle")
    subtitle.runs[0].bold = True
    metadata = doc.add_table(rows=2, cols=4)
    metadata.style = "Table Grid"
    values = [
        ("学科", info.get("subject") or "—"),
        ("年级", info.get("grade_level") or info.get("audience") or "—"),
        ("总分", f"{settings['total_score']} 分"),
        ("建议用时", f"{settings['estimated_minutes']} 分钟"),
    ]
    for row_index in range(2):
        for pair_index in range(2):
            label, value = values[row_index * 2 + pair_index]
            label_cell = metadata.cell(row_index, pair_index * 2)
            _shade_cell(label_cell, "E8EEF5")
            _write_cell(label_cell, label, bold=True, color="1F4D78")
            _write_cell(metadata.cell(row_index, pair_index * 2 + 1), value)
    _set_table_geometry(metadata, [1400, 3280, 1400, 3280])
    _add_heading(doc, "作答说明", 1)
    for instruction in settings["student_instructions"]:
        doc.add_paragraph(instruction, style="List Bullet")
    doc.add_paragraph(settings["answer_requirements"], style="List Bullet")

    question_number = 0
    for section in content["sections"]:
        _add_heading(doc, f"{section['title']}（{section['score']} 分）", 1)
        for block in section["blocks"]:
            if block["kind"] == "question_group":
                _add_heading(doc, f"{block['id']} · {block['title']}", 2)
                if block.get("instructions"):
                    doc.add_paragraph(block["instructions"])
                for stimulus in block["stimuli"]:
                    _add_exercise_stimulus(doc, stimulus, asset_paths)
                for item in block["sub_questions"]:
                    question_number += 1
                    _add_exercise_question(doc, item, question_number, include_answers)
            else:
                question_number += 1
                _add_exercise_question(doc, block, question_number, include_answers)
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)
    Document(output)
    return output


def _set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    marker = OxmlElement("w:tblHeader")
    marker.set(qn("w:val"), "true")
    tr_pr.append(marker)


def _shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def _set_cell_margins(cell, top: int = 80, bottom: int = 80, start: int = 120, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_table_geometry(table, widths: list[int], indent: int = 120) -> None:
    table.autofit = False
    total = sum(widths)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths[min(index, len(widths) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            _set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _set_style_font(style, name: str, east_asia: str, size: float, color: str | None = None) -> None:
    style.font.name = name
    style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia)
    style.font.size = Pt(size)
    if color:
        style.font.color.rgb = RGBColor.from_string(color)


def _configure_task_sheet(doc: Document, title: str, version: str, edition: str = "学生版") -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = section.right_margin = section.bottom_margin = section.left_margin = Inches(1)
    section.header_distance = section.footer_distance = Inches(0.492)
    header = section.header.paragraphs[0]
    header.text = f"LessonForge AI · {title} · {version}"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        run.font.name = TASK_SHEET_CJK_FONT
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor.from_string("656A73")
        run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), TASK_SHEET_CJK_FONT)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.add_run(f"{edition} · 第 ")
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    footer.runs[-1]._r.addnext(field)
    footer.add_run(" 页")
    for run in footer.runs:
        run.font.name = TASK_SHEET_CJK_FONT
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor.from_string("656A73")
        run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), TASK_SHEET_CJK_FONT)

    styles = doc.styles
    normal = styles["Normal"]
    _set_style_font(normal, TASK_SHEET_CJK_FONT, TASK_SHEET_CJK_FONT, 11, "18191D")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for name, size, color, before, after in (
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ):
        style = styles[name]
        _set_style_font(style, TASK_SHEET_CJK_FONT, TASK_SHEET_CJK_FONT, size, color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
    title_style = styles["Title"]
    _set_style_font(title_style, TASK_SHEET_CJK_FONT, TASK_SHEET_CJK_FONT, 24, "0B2545")
    title_style.font.bold = True
    title_style.paragraph_format.space_before = Pt(0)
    title_style.paragraph_format.space_after = Pt(4)
    subtitle_style = styles["Subtitle"]
    _set_style_font(subtitle_style, TASK_SHEET_CJK_FONT, TASK_SHEET_CJK_FONT, 11, "656A73")
    subtitle_style.paragraph_format.space_before = Pt(0)
    subtitle_style.paragraph_format.space_after = Pt(14)
    for name in ("List Bullet", "List Number"):
        style = styles[name]
        _set_style_font(style, TASK_SHEET_CJK_FONT, TASK_SHEET_CJK_FONT, 11, "18191D")
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25


def _add_heading(doc: Document, text: str, level: int) -> None:
    paragraph = doc.add_heading(text, level=level)
    paragraph.paragraph_format.keep_with_next = True


def _write_cell(cell, text: str, *, bold: bool = False, color: str = "18191D", align=None) -> None:
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = TASK_SHEET_CJK_FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), TASK_SHEET_CJK_FONT)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor.from_string(color)


def _render_record_table_docx(doc: Document, record: dict) -> None:
    _add_heading(doc, record["title"], 3)
    doc.add_paragraph(record["instructions"])
    columns = record["columns"]
    table = doc.add_table(rows=1 + record["blank_rows"], cols=len(columns))
    table.style = "Table Grid"
    for index, column in enumerate(columns):
        _shade_cell(table.cell(0, index), "E8EEF5")
        _write_cell(table.cell(0, index), column, bold=True, color="1F4D78", align=WD_ALIGN_PARAGRAPH.CENTER)
    _set_repeat_table_header(table.rows[0])
    for row in table.rows[1:]:
        row.height = Inches(0.42)
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    base = 9360 // len(columns)
    widths = [base] * len(columns)
    widths[-1] += 9360 - sum(widths)
    _set_table_geometry(table, widths)


def render_task_sheet_v3_docx(title: str, content: dict, output: Path, version: str = "V3") -> Path:
    """V3 动态目录任务单 DOCX：按目录树层级递归输出 Heading 1–3 与强类型 Block。"""
    doc = Document()
    _configure_task_sheet(doc, title, version)
    info = content["course_info"]

    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_after = Pt(4)
    title_run = title_p.add_run(info["course_title"])
    title_run.bold = True
    title_run.font.name = TASK_SHEET_CJK_FONT
    title_run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), TASK_SHEET_CJK_FONT)
    title_run.font.size = Pt(24)
    title_run.font.color.rgb = RGBColor.from_string("0B2545")
    subtitle = doc.add_paragraph("学习任务单 · 学生版 · 动态版")
    subtitle.paragraph_format.space_after = Pt(14)
    for run in subtitle.runs:
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor.from_string("656A73")

    metadata = doc.add_table(rows=2, cols=4)
    metadata.style = "Table Grid"
    metadata_values = [
        ("学科", info.get("subject") or "—"), ("年级", info.get("grade_level") or info.get("audience") or "—"),
        ("建议时长", f"{info['duration_minutes']} 分钟"), ("学习对象", info.get("audience") or "—"),
    ]
    for row_index in range(2):
        for pair_index in range(2):
            label, value = metadata_values[row_index * 2 + pair_index]
            label_cell = metadata.cell(row_index, pair_index * 2)
            value_cell = metadata.cell(row_index, pair_index * 2 + 1)
            _shade_cell(label_cell, "E8EEF5")
            _write_cell(label_cell, label, bold=True, color="1F4D78")
            _write_cell(value_cell, str(value))
    _set_table_geometry(metadata, [1500, 3180, 1500, 3180])

    catalog = {item["id"]: item for item in content.get("objective_catalog", [])}

    def _objective_texts(objective_ids: list[str]) -> list[str]:
        return [
            f"{objective_id} · {catalog[objective_id]['statement']}（达成标准：{catalog[objective_id]['success_criterion']}）"
            for objective_id in objective_ids
            if objective_id in catalog
        ]

    def render_block(block: dict) -> None:
        kind = block.get("kind")
        if kind == "text":
            doc.add_paragraph(block.get("text", ""))
        elif kind == "objective_list":
            _add_heading(doc, block.get("title") or "学习目标", 2)
            for text in _objective_texts(block.get("objective_ids", [])):
                paragraph = doc.add_paragraph(style="List Bullet")
                paragraph.add_run(text)
        elif kind == "learning_task":
            collaboration_labels = {"individual": "独立", "pair": "结对", "group": "小组", "whole_class": "全班"}
            _add_heading(doc, f"{block['id']} · {block['title']}", 2)
            meta = doc.add_table(rows=2, cols=4)
            meta.style = "Table Grid"
            values = [
                ("预计用时", f"{block.get('estimated_minutes')} 分钟"),
                ("协作方式", collaboration_labels.get(block.get("collaboration_mode", ""), block.get("collaboration_mode", "独立"))),
                ("对应目标", "、".join(block.get("objective_ids", []))),
                ("教学环节", block.get("stage_id") or "—"),
            ]
            for row_index in range(2):
                for pair_index in range(2):
                    label, value = values[row_index * 2 + pair_index]
                    label_cell = meta.cell(row_index, pair_index * 2)
                    _shade_cell(label_cell, "F2F4F7")
                    _write_cell(label_cell, label, bold=True, color="1F4D78")
                    _write_cell(meta.cell(row_index, pair_index * 2 + 1), str(value))
            _set_table_geometry(meta, [1400, 3280, 1400, 3280])
            doc.add_paragraph(f"学习动作：{block.get('action')}；操作对象：{block.get('object')}")
            _add_heading(doc, "操作步骤", 3)
            for step in block.get("steps", []):
                doc.add_paragraph(step, style="List Number")
            output_p = doc.add_paragraph()
            output_p.add_run("成果要求：").bold = True
            output_p.add_run(block.get("student_output", ""))
            criterion_p = doc.add_paragraph()
            criterion_p.add_run("完成标准：").bold = True
            criterion_p.add_run(block.get("completion_criterion", ""))
            for scaffold in block.get("scaffolds", []):
                doc.add_paragraph(scaffold, style="List Bullet")
            if block.get("record_table"):
                _render_record_table_docx(doc, block["record_table"])
        elif kind == "record_table":
            _render_record_table_docx(doc, block)
        elif kind == "question_set":
            _add_heading(doc, block.get("title") or "课堂问题", 2)
            for question in block.get("questions", []):
                paragraph = doc.add_paragraph()
                paragraph.add_run(f"{question['id']} · ").bold = True
                paragraph.add_run(question["prompt"])
                response = doc.add_table(rows=2, cols=1)
                response.style = "Table Grid"
                for row in response.rows:
                    row.height = Inches(0.35)
                    row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
                _set_table_geometry(response, [9360])
        elif kind == "assessment":
            scale = block.get("scale", ["尚未做到", "基本做到", "能够做到"])
            items = block.get("items", [])
            assessment = doc.add_table(rows=1 + len(items), cols=1 + len(scale))
            assessment.style = "Table Grid"
            headers = ["自评项目", *scale]
            for index, header in enumerate(headers):
                _shade_cell(assessment.cell(0, index), "E8EEF5")
                _write_cell(assessment.cell(0, index), header, bold=True, color="1F4D78", align=WD_ALIGN_PARAGRAPH.CENTER)
            for row_index, item in enumerate(items, 1):
                _write_cell(assessment.cell(row_index, 0), item["statement"])
                for column_index in range(1, len(headers)):
                    _write_cell(assessment.cell(row_index, column_index), "", align=WD_ALIGN_PARAGRAPH.CENTER)
                assessment.rows[row_index].height = Inches(0.3)
                assessment.rows[row_index].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
            label_width = 4800
            scale_width = (9360 - label_width) // len(scale)
            widths = [label_width, *([scale_width] * len(scale))]
            widths[-1] += 9360 - sum(widths)
            _set_repeat_table_header(assessment.rows[0])
            _set_table_geometry(assessment, widths)
        elif kind == "checklist":
            _add_heading(doc, block.get("title") or "检查表", 2)
            for item in block.get("items", []):
                doc.add_paragraph(f"□ {item['text']}", style="List Bullet")

    # 深度优先遍历（parent_id/order 表达目录树）
    sections = content.get("sections", [])
    depth_map: dict[str, int] = {}
    for section in sections:
        depth_map[section["id"]] = 1 if not section.get("parent_id") else depth_map.get(section.get("parent_id"), 1) + 1
    for section in sorted(sections, key=lambda s: (s.get("parent_id") or "", int(s.get("order", 0)))):
        level = min(depth_map.get(section["id"], 1) + 1, 3)
        _add_heading(doc, section["title"], level)
        if section.get("purpose"):
            purpose_p = doc.add_paragraph()
            purpose_run = purpose_p.add_run(section["purpose"])
            purpose_run.italic = True
            purpose_run.font.color.rgb = RGBColor.from_string("656A73")
        for block in section.get("blocks", []):
            render_block(block)

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)
    Document(output)
    return output


def render_task_sheet_docx(title: str, content: dict, output: Path, version: str = "V1") -> Path:
    if content.get("schema_version") == "3.0":
        return render_task_sheet_v3_docx(title, content, output, "V3")
    doc = Document()
    _configure_task_sheet(doc, title, version)
    info = content["course_info"]

    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_after = Pt(4)
    title_run = title_p.add_run(info["course_title"])
    title_run.bold = True
    title_run.font.name = TASK_SHEET_CJK_FONT
    title_run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), TASK_SHEET_CJK_FONT)
    title_run.font.size = Pt(24)
    title_run.font.color.rgb = RGBColor.from_string("0B2545")
    subtitle = doc.add_paragraph("学习任务单 · 学生版")
    subtitle.paragraph_format.space_after = Pt(14)
    for run in subtitle.runs:
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor.from_string("656A73")

    metadata = doc.add_table(rows=2, cols=4)
    metadata.style = "Table Grid"
    metadata_values = [
        ("学科", info.get("subject") or "—"), ("年级", info.get("grade_level") or info.get("audience") or "—"),
        ("建议时长", f"{info['duration_minutes']} 分钟"), ("学习对象", info.get("audience") or "—"),
    ]
    for row_index in range(2):
        for pair_index in range(2):
            label, value = metadata_values[row_index * 2 + pair_index]
            label_cell = metadata.cell(row_index, pair_index * 2)
            value_cell = metadata.cell(row_index, pair_index * 2 + 1)
            _shade_cell(label_cell, "E8EEF5")
            _write_cell(label_cell, label, bold=True, color="1F4D78")
            _write_cell(value_cell, str(value))
    _set_table_geometry(metadata, [1500, 3180, 1500, 3180])

    _add_heading(doc, "学习目标", 1)
    for objective in content["learning_objectives"]:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.add_run(f"{objective['id']} · {objective['statement']} ").bold = True
        paragraph.add_run(f"达成标准：{objective['success_criterion']}")

    _add_heading(doc, "课前准备", 1)
    for item in content.get("preparation", []):
        doc.add_paragraph(item, style="List Bullet")

    phase_labels = {"pre_class": "课前", "in_class": "课中", "after_class": "课后"}
    collaboration_labels = {"individual": "独立", "pair": "结对", "group": "小组", "whole_class": "全班"}
    _add_heading(doc, "学习任务", 1)
    for task in content["tasks"]:
        _add_heading(doc, f"{task['id']} · {task['title']}", 2)
        meta = doc.add_table(rows=2, cols=4)
        meta.style = "Table Grid"
        values = [
            ("阶段", phase_labels.get(task["phase"], task["phase"])),
            ("预计用时", f"{task['estimated_minutes']} 分钟"),
            ("协作方式", collaboration_labels.get(task["collaboration_mode"], task["collaboration_mode"])),
            ("对应目标", "、".join(task["objective_ids"])),
        ]
        for row_index in range(2):
            for pair_index in range(2):
                label, value = values[row_index * 2 + pair_index]
                label_cell = meta.cell(row_index, pair_index * 2)
                _shade_cell(label_cell, "F2F4F7")
                _write_cell(label_cell, label, bold=True, color="1F4D78")
                _write_cell(meta.cell(row_index, pair_index * 2 + 1), str(value))
        _set_table_geometry(meta, [1400, 3280, 1400, 3280])
        doc.add_paragraph(f"学习动作：{task['action']}；操作对象：{task['object']}")
        _add_heading(doc, "操作步骤", 3)
        for step in task["steps"]:
            doc.add_paragraph(step, style="List Number")
        output_p = doc.add_paragraph()
        output_p.add_run("成果要求：").bold = True
        output_p.add_run(task["student_output"])
        criterion_p = doc.add_paragraph()
        criterion_p.add_run("完成标准：").bold = True
        criterion_p.add_run(task["completion_criterion"])
        if task.get("scaffolds"):
            _add_heading(doc, "思考支架", 3)
            for scaffold in task["scaffolds"]:
                doc.add_paragraph(scaffold, style="List Bullet")
        record = task.get("record_table")
        if record:
            _add_heading(doc, record["title"], 3)
            doc.add_paragraph(record["instructions"])
            columns = record["columns"]
            table = doc.add_table(rows=1 + record["blank_rows"], cols=len(columns))
            table.style = "Table Grid"
            for index, column in enumerate(columns):
                _shade_cell(table.cell(0, index), "E8EEF5")
                _write_cell(table.cell(0, index), column, bold=True, color="1F4D78", align=WD_ALIGN_PARAGRAPH.CENTER)
            _set_repeat_table_header(table.rows[0])
            for row in table.rows[1:]:
                row.height = Inches(0.42)
                row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
            base = 9360 // len(columns)
            widths = [base] * len(columns)
            widths[-1] += 9360 - sum(widths)
            _set_table_geometry(table, widths)

    global_record = content.get("record_table")
    if global_record:
        _add_heading(doc, global_record["title"], 1)
        doc.add_paragraph(global_record["instructions"])
        columns = global_record["columns"]
        table = doc.add_table(rows=1 + global_record["blank_rows"], cols=len(columns))
        table.style = "Table Grid"
        for index, column in enumerate(columns):
            _shade_cell(table.cell(0, index), "E8EEF5")
            _write_cell(table.cell(0, index), column, bold=True, color="1F4D78", align=WD_ALIGN_PARAGRAPH.CENTER)
        _set_repeat_table_header(table.rows[0])
        for row in table.rows[1:]:
            row.height = Inches(0.42)
            row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        base = 9360 // len(columns)
        widths = [base] * len(columns)
        widths[-1] += 9360 - sum(widths)
        _set_table_geometry(table, widths)

    _add_heading(doc, "课堂问题", 1)
    for question in content.get("learning_questions", []):
        paragraph = doc.add_paragraph()
        paragraph.add_run(f"{question['id']} · ").bold = True
        paragraph.add_run(question["prompt"])
        response = doc.add_table(rows=2, cols=1)
        response.style = "Table Grid"
        for row in response.rows:
            row.height = Inches(0.35)
            row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        _set_table_geometry(response, [9360])

    doc.add_page_break()
    _add_heading(doc, "自我评价", 1)
    scale = content["self_assessment_scale"]
    assessment = doc.add_table(rows=1 + len(content["self_assessment"]), cols=1 + len(scale))
    assessment.style = "Table Grid"
    headers = ["自评项目", *scale]
    for index, header in enumerate(headers):
        _shade_cell(assessment.cell(0, index), "E8EEF5")
        _write_cell(assessment.cell(0, index), header, bold=True, color="1F4D78", align=WD_ALIGN_PARAGRAPH.CENTER)
    for row_index, item in enumerate(content["self_assessment"], 1):
        _write_cell(assessment.cell(row_index, 0), item["statement"])
        for column_index in range(1, len(headers)):
            _write_cell(assessment.cell(row_index, column_index), "", align=WD_ALIGN_PARAGRAPH.CENTER)
        assessment.rows[row_index].height = Inches(0.3)
        assessment.rows[row_index].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    label_width = 4800
    scale_width = (9360 - label_width) // len(scale)
    widths = [label_width, *([scale_width] * len(scale))]
    widths[-1] += 9360 - sum(widths)
    _set_repeat_table_header(assessment.rows[0])
    _set_table_geometry(assessment, widths)

    _add_heading(doc, "课后拓展", 1)
    for item in content.get("extension", []):
        doc.add_paragraph(item, style="List Bullet")
    extension_space = doc.add_table(rows=3, cols=1)
    extension_space.style = "Table Grid"
    for row in extension_space.rows:
        row.height = Inches(0.35)
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    _set_table_geometry(extension_space, [9360])

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)
    Document(output)
    return output


def render_video_script_docx(
    title: str,
    content: dict,
    output: Path,
    version: str = "V2",
    source_versions: dict[str, int] | None = None,
) -> Path:
    doc = Document()
    _configure(doc, title, version)
    info = content["course_info"]
    settings = content["production_settings"]
    source_versions = source_versions or {}

    doc.add_heading(info["course_title"], 0)
    doc.add_paragraph("微课视频脚本 · 录屏制作版", style="Subtitle")
    metadata = doc.add_table(rows=3, cols=4)
    metadata.style = "Table Grid"
    values = [
        ("学科", info.get("subject") or "—"),
        ("年级 / 对象", info.get("grade_level") or info.get("audience") or "—"),
        ("制作方式", "16:9 PPT 录屏与常规动效"),
        ("目标时长", f"{settings['target_duration_seconds']} 秒"),
        ("建议语速", f"{settings['narration_chars_per_minute']} 字/分钟"),
        ("上游版本", f"教学设计 V{source_versions.get('lesson_plan', '—')} · PPT V{source_versions.get('ppt', '—')}"),
    ]
    for row_index in range(3):
        for pair_index in range(2):
            label, value = values[row_index * 2 + pair_index]
            label_cell = metadata.cell(row_index, pair_index * 2)
            _shade_cell(label_cell, "E8EEF5")
            _write_cell(label_cell, label, bold=True, color="1F4D78")
            _write_cell(metadata.cell(row_index, pair_index * 2 + 1), str(value))
    _set_table_geometry(metadata, [1400, 3280, 1400, 3280])

    for scene in content["scenes"]:
        doc.add_page_break()
        _add_heading(doc, f"{scene['id']} · {scene['title']}", 1)
        meta = doc.add_table(rows=2, cols=4)
        meta.style = "Table Grid"
        scene_values = [
            ("时间", f"{scene['start_seconds']:.0f}s—{scene['end_seconds']:.0f}s"),
            ("教学角色", scene["pedagogical_role"]),
            ("PPT / 环节", f"{scene['slide_id']} / {scene['lesson_stage_id']}"),
            ("目标 / 知识点", f"{'、'.join(scene['objective_ids'])} / {'、'.join(scene['knowledge_point_ids'])}"),
        ]
        for row_index in range(2):
            for pair_index in range(2):
                label, value = scene_values[row_index * 2 + pair_index]
                label_cell = meta.cell(row_index, pair_index * 2)
                _shade_cell(label_cell, "F2F4F7")
                _write_cell(label_cell, label, bold=True, color="1F4D78")
                _write_cell(meta.cell(row_index, pair_index * 2 + 1), str(value))
        _set_table_geometry(meta, [1400, 3280, 1400, 3280])

        visual = scene["visual_track"]
        audio = scene["audio_track"]
        text_track = scene["text_track"]
        animation = "\n".join(
            f"+{cue['offset_seconds']}s · {cue['action']} {cue['target']}：{cue['instruction']}"
            for cue in visual.get("animation_cues", [])
        ) or "无"
        pauses = "；".join(
            f"+{cue['offset_seconds']}s 停顿 {cue['duration_seconds']}s（{cue['purpose']}）"
            for cue in audio.get("pause_cues", [])
        ) or "无"
        sounds = "；".join(
            f"+{cue['offset_seconds']}s {cue['description']}" for cue in audio.get("sound_cues", [])
        ) or "无"
        subtitles = "\n".join(
            f"+{cue['start_offset_seconds']}s—{cue['end_offset_seconds']}s · {cue['text']}"
            for cue in text_track.get("subtitle_chunks", [])
        )
        interaction = scene.get("interaction")
        interaction_text = "无"
        if interaction:
            interaction_text = (
                f"问题：{interaction['prompt']}\n等待：{interaction['wait_seconds']} 秒\n"
                f"预期回应：{interaction['expected_response']}\n反馈衔接：{interaction['feedback_transition']}"
            )
        production = doc.add_table(rows=7, cols=2)
        production.style = "Table Grid"
        rows = [
            ("学习目的", scene["learning_purpose"]),
            ("画面与构图", visual["composition"]),
            ("动效与转场", animation),
            ("旁白与声音", f"{audio['narration_text']}\n\n语气：{audio['delivery_tone']}\n强调：{'、'.join(audio['emphasis_terms']) or '无'}\n停顿：{pauses}\n音效：{sounds}"),
            ("字幕与屏显", f"{subtitles}\n\n屏幕贴字：{'、'.join(text_track['on_screen_text']) or '无'}"),
            ("互动", interaction_text),
            ("制作备注", "；".join(scene.get("production_notes", [])) or "无"),
        ]
        for row_index, (label, value) in enumerate(rows):
            _shade_cell(production.cell(row_index, 0), "E8EEF5")
            _write_cell(production.cell(row_index, 0), label, bold=True, color="1F4D78")
            _write_cell(production.cell(row_index, 1), value)
        _set_table_geometry(production, [1800, 7560])

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)
    Document(output)
    return output
