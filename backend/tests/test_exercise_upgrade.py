from pathlib import Path

import pytest
from docx import Document
from fastapi import HTTPException
from pydantic import ValidationError

from app.agents.generators import make_blueprint, make_exercises, to_markdown
from app.api.v1.settings import _validate_model_transport
from app.models.entities import CourseProject
from app.renderers.docx_renderer import render_exercise_docx
from app.schemas.artifact import ExerciseContent
from app.services.exercise_review_service import degrade_unreviewed_visuals
from app.services.exercise_visual_service import _safe_remote_image, render_deterministic_svg
from app.services.quality_service import validate_exercise


def course(duration: int = 20):
    return CourseProject(
        owner_id="teacher", title="阿基米德原理", subject="物理", grade_level="八年级",
        audience="已学习液体压强的学生", duration_minutes=duration, scenario="课堂讲解",
        language="中文", settings_json={"course_task": "解释浮力并分析新情境"},
    )


def document_text(document: Document) -> str:
    values = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                values.extend(paragraph.text for paragraph in cell.paragraphs)
    return "\n".join(values)


def test_exercise_v2_schema_quality_and_student_markdown():
    blueprint = make_blueprint(course())
    exercise = make_exercises(blueprint)
    validated = ExerciseContent.model_validate(exercise.model_dump())

    assert validated.schema_version == "2.0"
    assert [section.id for section in validated.sections] == [
        "basic_consolidation", "understanding_application", "transfer_challenge",
    ]
    assert sum(section.score for section in validated.sections) == 100
    assert validate_exercise(blueprint, validated.model_dump()) == []

    markdown = to_markdown("exercise", validated)
    assert "基础巩固" in markdown and "理解应用" in markdown and "迁移挑战" in markdown
    assert "参考答案" not in markdown and "**答案：**" not in markdown and "**解析：**" not in markdown


def test_exercise_v2_rejects_score_and_rubric_drift():
    blueprint = make_blueprint(course())
    payload = make_exercises(blueprint).model_dump()
    payload["sections"][2]["blocks"][0]["scoring_points"][0]["points"] = 1
    with pytest.raises(ValidationError):
        ExerciseContent.model_validate(payload)

    invalid_ref = make_exercises(blueprint).model_dump()
    invalid_ref["sections"][0]["blocks"][0]["objective_ids"] = ["OBJ-404"]
    issues = validate_exercise(blueprint, invalid_ref)
    assert any("OBJ-404" in item["description"] for item in issues)
    assert all(item["target_agent"] == "exercise_agent" for item in issues)
    assert all(item["location"].startswith("$") for item in issues)


def test_student_and_teacher_docx_are_separated(tmp_path: Path):
    exercise = make_exercises(make_blueprint(course())).model_dump()
    student_path = render_exercise_docx("学生卷", exercise, tmp_path / "student.docx", False)
    teacher_path = render_exercise_docx("教师卷", exercise, tmp_path / "teacher.docx", True)
    student_text = document_text(Document(student_path))
    teacher_text = document_text(Document(teacher_path))

    assert "参考答案" not in student_text and "评分点" not in student_text and "常见错误" not in student_text
    assert "参考答案" in teacher_text and "评分点" in teacher_text and "常见错误" in teacher_text
    assert len(Document(student_path).tables) >= 1
    assert len(teacher_text) > len(student_text)


def test_unreviewed_image_degrades_and_deterministic_svg_is_stable():
    payload = make_exercises(make_blueprint(course())).model_dump()
    group = payload["sections"][1]["blocks"][0]
    group["stimuli"] = [{
        "id": "VIS-ST-01", "kind": "visual", "title": "实验现象", "text": "", "columns": [], "rows": [],
        "visual": {
            "visual_id": "VIS-01", "mode": "generated_image", "purpose": "观察液面现象",
            "alt_text": "装有液体的透明容器", "caption": "", "fallback_stimulus": "容器内液面保持水平。",
            "generation_prompt": "透明容器中的液体现象，高对比教学图片", "size": "1536x1024",
            "diagram_type": None, "diagram_spec": {}, "asset_id": None, "status": "requested",
            "provider": "", "model_name": "", "review_notes": [],
        },
    }]
    degraded, notes = degrade_unreviewed_visuals(payload)
    stimulus = degraded["sections"][1]["blocks"][0]["stimuli"][0]
    assert stimulus["kind"] == "text"
    assert "液面保持水平" in stimulus["text"]
    assert notes and "自动使用等价文字材料" in notes[0]

    first = render_deterministic_svg("flow", {"nodes": [{"label": "条件"}, {"label": "检查"}]}, "处理流程")
    second = render_deterministic_svg("flow", {"nodes": [{"label": "条件"}, {"label": "检查"}]}, "处理流程")
    assert first == second
    assert first.startswith(b"<svg") and "处理流程" in first.decode("utf-8")


@pytest.mark.asyncio
async def test_visual_transport_rejects_insecure_or_executable_configuration():
    with pytest.raises(ValueError, match="HTTPS"):
        await _safe_remote_image("http://127.0.0.1/internal.png", 1)

    with pytest.raises(HTTPException, match="站内绝对路径"):
        _validate_model_transport(
            "openai_compatible", "custom_image_http", "https://images.example.com",
            {"endpoint_path": "https://evil.example.com/generate"},
        )
    with pytest.raises(HTTPException, match="不受支持的字段"):
        _validate_model_transport(
            "openai_compatible", "custom_image_http", "https://images.example.com",
            {"endpoint_path": "/generate", "script": "return response"},
        )

    _validate_model_transport(
        "openai_compatible", "custom_image_http", "https://images.example.com",
        {
            "endpoint_path": "/generate", "auth_mode": "x_api_key",
            "prompt_field": "prompt", "response_base64_path": "data.image",
        },
    )
