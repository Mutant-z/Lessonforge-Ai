from pathlib import Path
from types import SimpleNamespace

import pytest
from docx import Document
from pydantic import ValidationError

from app.agents.generators import (
    make_blueprint,
    make_lesson_plan,
    make_ppt,
    make_verbatim,
    make_video_script,
    recalculate_scene_timelines,
    repair_video_script_subtitles,
    to_markdown,
)
from app.models.entities import CourseProject
from app.renderers.docx_renderer import render_video_script_docx
from app.schemas.artifact import (
    AnimationCue,
    LegacyVideoScriptContent,
    PauseCue,
    SubtitleChunk,
    VideoAudioTrack,
    VideoScene,
    VideoScriptContent,
    VideoTextTrack,
    VideoVisualTrack,
)
from app.services.quality_service import validate_video_script
from app.services.video_generation_service import (
    _ai_visual_scene_ids,
    _compose,
    _local_tts_engine,
    _generate_local_speech,
    _hydrate_scene_director_data,
    _media_error_detail,
    _pcm_rms,
    _retryable_image_error,
    _shot_plan,
    _subtitle_documents,
    _validate_sources,
    _visual_prompt,
)
from app.schemas.video import VideoGenerationContent, VideoGenerationScene


def sample_course():
    return CourseProject(
        owner_id="u", title="牛顿第二定律", subject="高中物理", grade_level="高一",
        audience="已学习运动学基础的学生", duration_minutes=15, scenario="课堂讲解",
        language="中文", settings_json={"course_task": "解释力、质量与加速度的关系"},
    )


def test_video_script_v2_schema_quality_and_verbatim_handoff():
    bp = make_blueprint(sample_course())
    lesson, ppt = make_lesson_plan(bp), make_ppt(bp)
    script = make_video_script(bp, lesson, ppt)

    assert script.schema_version == "2.0"
    assert script.production_settings.mode == "ppt_screen_recording"
    assert script.scenes[-1].end_seconds == bp.course_identity.duration_minutes * 60
    assert {scene.slide_id for scene in script.scenes} == {slide.id for slide in ppt.slides}
    assert validate_video_script(bp, script.model_dump(), lesson.model_dump(), ppt.model_dump()) == []

    verbatim = make_verbatim(bp, ppt, script)
    assert len(verbatim.sections) == len(script.scenes)
    assert verbatim.sections[0].slide_ids == [script.scenes[0].slide_id]
    assert script.scenes[0].audio_track.narration_text in verbatim.sections[0].required_text


def test_mock_video_script_preserves_edited_ppt_timing_as_production_target():
    bp = make_blueprint(sample_course())
    lesson, ppt = make_lesson_plan(bp), make_ppt(bp)
    ppt.slides[0].duration_seconds += 37

    script = make_video_script(bp, lesson, ppt)

    assert script.scenes[0].start_seconds == 0
    assert script.scenes[-1].end_seconds == sum(slide.duration_seconds for slide in ppt.slides)
    assert script.production_settings.target_duration_seconds == script.scenes[-1].end_seconds
    assert all(
        current.start_seconds == previous.end_seconds
        for previous, current in zip(script.scenes, script.scenes[1:])
    )


def test_video_script_schema_rejects_broken_timeline_and_cues():
    bp = make_blueprint(sample_course())
    lesson, ppt = make_lesson_plan(bp), make_ppt(bp)
    payload = make_video_script(bp, lesson, ppt).model_dump()
    payload["scenes"][1]["start_seconds"] = payload["scenes"][0]["start_seconds"]
    with pytest.raises(ValidationError, match="重叠|倒序"):
        VideoScriptContent.model_validate(payload)

    payload = make_video_script(bp, lesson, ppt).model_dump()
    payload["scenes"][0]["visual_track"]["animation_cues"][0]["offset_seconds"] = 999
    with pytest.raises(ValidationError, match="动效提示超出"):
        VideoScriptContent.model_validate(payload)


def test_video_script_quality_reports_external_reference_and_subtitle_errors():
    bp = make_blueprint(sample_course())
    lesson, ppt = make_lesson_plan(bp), make_ppt(bp)
    payload = make_video_script(bp, lesson, ppt).model_dump()
    payload["scenes"][0]["slide_id"] = "S404"
    payload["scenes"][0]["lesson_stage_id"] = "ACT-404"
    payload["scenes"][0]["objective_ids"] = ["OBJ-404"]
    payload["scenes"][0]["knowledge_point_ids"] = ["KP-404"]
    payload["scenes"][1]["text_track"]["subtitle_chunks"][0]["text"] = "字幕不一致"

    issues = validate_video_script(bp, payload, lesson.model_dump(), ppt.model_dump())
    descriptions = "\n".join(item["description"] for item in issues)
    assert all(item["target_agent"] == "video_script_agent" for item in issues)
    assert "S404" in descriptions
    assert "ACT-404" in descriptions
    assert "OBJ-404" in descriptions
    assert "KP-404" in descriptions
    assert "字幕未完整" in descriptions


def test_video_script_subtitle_repair_rebuilds_only_mismatched_tracks():
    bp = make_blueprint(sample_course())
    lesson, ppt = make_lesson_plan(bp), make_ppt(bp)
    script = make_video_script(bp, lesson, ppt)
    original_second_scene = script.scenes[1].text_track.model_dump()
    script.scenes[0].text_track.subtitle_chunks[0].text = "字幕不一致"

    repaired = repair_video_script_subtitles(script)

    assert "".join(
        chunk.text for chunk in repaired.scenes[0].text_track.subtitle_chunks
    ) == repaired.scenes[0].audio_track.narration_text
    assert repaired.scenes[1].text_track.model_dump() == original_second_scene
    assert script.scenes[0].text_track.subtitle_chunks[0].text == "字幕不一致"
    issues = validate_video_script(bp, repaired.model_dump(), lesson.model_dump(), ppt.model_dump())
    assert not any("字幕未完整" in item["description"] for item in issues)


def test_legacy_video_script_remains_readable_and_requests_v2_upgrade():
    bp = make_blueprint(sample_course())
    lesson, ppt = make_lesson_plan(bp), make_ppt(bp)
    legacy = {
        "segments": [{
            "id": "VS-01", "time_range": "00:00—00:20", "stage": "导入",
            "slide_ids": ["S01"], "visual": "显示封面", "narration": "欢迎学习",
            "action": "淡入", "on_screen_text": "课程标题", "pause": "自然停顿",
            "production_notes": "保持 16:9",
        }],
    }
    assert LegacyVideoScriptContent.model_validate(legacy).segments[0].id == "VS-01"
    issues = validate_video_script(bp, legacy, lesson.model_dump(), ppt.model_dump())
    assert issues[0]["dimension"] == "compatibility"


@pytest.mark.asyncio
async def test_video_generation_preflight_returns_actionable_error_for_legacy_script():
    legacy = {
        "segments": [{
            "id": "VS-01", "time_range": "00:00—00:20", "stage": "导入",
            "slide_ids": ["S01"], "visual": "显示封面", "narration": "欢迎学习",
            "action": "淡入", "on_screen_text": "课程标题", "pause": "自然停顿",
            "production_notes": "保持 16:9",
        }],
    }

    class FakeDb:
        def __init__(self):
            self.results = iter([
                SimpleNamespace(id="script-1", content_json=legacy),
                SimpleNamespace(id="ppt-1", content_json={}),
                SimpleNamespace(id="lesson-1", content_json={}),
                SimpleNamespace(id="blueprint-1", content_json={}),
                SimpleNamespace(status="review"),
            ])

        async def get(self, _model, _id):
            return SimpleNamespace(id="course-1", current_blueprint_version=1)

        async def scalar(self, _query):
            return next(self.results)

    with pytest.raises(ValueError, match="旧版结构.*同步最新内容"):
        await _validate_sources(FakeDb(), SimpleNamespace(course_id="course-1"))


def test_video_script_markdown_and_docx_cover_production_tracks(tmp_path: Path):
    bp = make_blueprint(sample_course())
    lesson, ppt = make_lesson_plan(bp), make_ppt(bp)
    script = make_video_script(bp, lesson, ppt)
    markdown = to_markdown("video_script", script)
    for marker in ("制作方式", "学习目的", "画面", "动效", "旁白", "字幕", "教学环节"):
        assert marker in markdown

    path = render_video_script_docx(
        "微课视频脚本", script.model_dump(), tmp_path / "video-script.docx", "V2",
        {"lesson_plan": 2, "ppt": 3},
    )
    document = Document(path)
    assert len(document.tables) >= 1 + len(script.scenes) * 2
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert "微课视频脚本" in text
    assert script.scenes[0].title in text


def test_script_v1_prompt_file_contains_seedance_v3_contract():
    prompt_path = Path(__file__).resolve().parents[1] / "app" / "prompts" / "script" / "v1.md"
    prompt = prompt_path.read_text(encoding="utf-8")
    assert len(prompt) > 0

    for field in (
        "continuity_group", "visual_prompt", "camera_beats", "spoken_text",
        "required_terms", "required_numbers", "required_facts", "voice_direction",
        "sound_design", "negative_constraints", "start_seconds", "end_seconds",
    ):
        assert field in prompt
    assert "Doubao-Seedance-2.5" in prompt
    assert "不得输出 `slide_id`" in prompt
    assert "豆包 ASR" in prompt


def _scene(
    scene_id: str,
    narration: str,
    pause_duration: float = 0.0,
    wait_seconds: float = 0.0,
) -> VideoScene:
    return VideoScene(
        id=scene_id,
        sequence=1,
        title="测试分镜",
        pedagogical_role="概念讲解",
        lesson_stage_id="ACT-01",
        slide_id="S01",
        objective_ids=["OBJ-01"],
        knowledge_point_ids=["KP-01"],
        start_seconds=0.0,
        end_seconds=1000.0,
        learning_purpose="测试",
        visual_track=VideoVisualTrack(
            composition="主画面",
            animation_cues=[AnimationCue(offset_seconds=999, target="页", action="高亮", instruction="测试")],
        ),
        audio_track=VideoAudioTrack(
            narration_text=narration,
            pedagogical_action="hook",
            delivery_tone="生动",
            pause_cues=[PauseCue(offset_seconds=0, duration_seconds=pause_duration, purpose="思考")] if pause_duration else [],
        ),
        text_track=VideoTextTrack(
            on_screen_text=["测试"],
            subtitle_chunks=[
                SubtitleChunk(start_offset_seconds=0, end_offset_seconds=30, text=narration),
            ],
        ),
        interaction=None if not wait_seconds else type(
            "Interaction", (), {"prompt": "请思考", "wait_seconds": wait_seconds,
                                "expected_response": "说出依据", "feedback_transition": "继续" })(),
    )


def test_recalculate_scene_timelines_recomputes_continuous_axis():
    # 第一镜旁白 6 字 @4cps => 1.5s，不足下限取 3.0s；第二镜 11 字 => 2.75s 口播 + 2.0s 停顿 => 4.75s
    scenes = [
        _scene("VS-01", "欢迎来到本课。", pause_duration=0.0),
        _scene("VS-02", "今天我们学习核心概念。", pause_duration=2.0),
    ]
    updated = recalculate_scene_timelines(scenes, chars_per_minute=240)

    assert updated[0].start_seconds == 0.0
    assert updated[0].end_seconds == 3.0
    assert updated[1].start_seconds == 3.0
    assert updated[1].end_seconds == 7.8  # 3.0 + 2.75s 口播 + 2.0s 停顿
    # 越界的动效与停顿 Cue 被裁剪到新时长内
    assert all(cue.offset_seconds <= 3.0 for cue in updated[0].visual_track.animation_cues)
    assert updated[1].audio_track.pause_cues and all(
        cue.offset_seconds + cue.duration_seconds <= 4.75 for cue in updated[1].audio_track.pause_cues
    )
    # 字幕块被重新对齐到新时长内
    assert updated[0].text_track.subtitle_chunks[-1].end_offset_seconds == 3.0
    assert updated[1].text_track.subtitle_chunks[-1].end_offset_seconds == 4.75


def test_make_video_script_populates_pedagogical_fields():
    bp = make_blueprint(sample_course())
    lesson, ppt = make_lesson_plan(bp), make_ppt(bp)
    script = make_video_script(bp, lesson, ppt)

    assert script.scenes[0].audio_track.pedagogical_action == "hook"
    assert script.scenes[-1].audio_track.pedagogical_action == "summary_recap"
    assert all(scene.audio_track.speaking_rate_cps == 4.0 for scene in script.scenes)
    assert validate_video_script(bp, script.model_dump(), lesson.model_dump(), ppt.model_dump()) == []


def test_make_verbatim_populates_pedagogical_and_timing_metadata():
    bp = make_blueprint(sample_course())
    lesson, ppt = make_lesson_plan(bp), make_ppt(bp)
    script = make_video_script(bp, lesson, ppt)
    verbatim = make_verbatim(bp, ppt, script)

    first = verbatim.sections[0]
    assert first.scene_id == script.scenes[0].id
    assert first.pedagogical_action == script.scenes[0].audio_track.pedagogical_action
    assert first.key_emphasis == script.scenes[0].audio_track.emphasis_terms
    narration_len = len(script.scenes[0].audio_track.narration_text)
    assert first.word_count == narration_len
    assert first.estimated_duration_seconds == round(narration_len / 4.0, 1)
    assert len(verbatim.sections) == len(script.scenes)


@pytest.mark.asyncio
async def test_video_compose_writes_absolute_concat_paths(tmp_path, monkeypatch):
    output_dir = tmp_path / "relative-storage" / "run"
    output_dir.mkdir(parents=True)
    scene_paths = [output_dir / "VG-01.mp4", output_dir / "VG-02.mp4"]
    for path in scene_paths:
        path.write_bytes(b"scene")
    srt_path = output_dir / "subtitles.srt"
    srt_path.write_text("", encoding="utf-8")

    async def fake_ffmpeg(*args):
        Path(args[-1]).write_bytes(b"video")

    async def fake_thumbnail(_video, output):
        output.write_bytes(b"thumbnail")

    async def fake_validate(_path):
        return None

    monkeypatch.setattr("app.services.video_generation_service._run_ffmpeg", fake_ffmpeg)
    monkeypatch.setattr("app.services.video_generation_service._thumbnail", fake_thumbnail)
    monkeypatch.setattr("app.services.video_generation_service._validate_media", fake_validate)

    await _compose(scene_paths, output_dir, "640x360", False, srt_path)

    concat_entries = (output_dir / "concat.txt").read_text(encoding="utf-8").splitlines()
    assert concat_entries == [f"file '{path.resolve().as_posix()}'" for path in scene_paths]


def _generation_scene(index: int, role: str) -> VideoGenerationScene:
    return VideoGenerationScene(
        id=f"VG-{index:02d}", script_scene_id=f"VS-{index:02d}", sequence=index,
        start_seconds=(index - 1) * 10, end_seconds=index * 10,
        visual_prompt="物理实验画面", narration_text="讲解内容", pedagogical_role=role,
        source_slide_id=f"slide_{index:02d}",
    )


def test_hybrid_director_selects_at_most_eight_key_visuals_and_builds_shots():
    roles = ["情境", "示范", "概念讲解", "导入", "检查点"] * 2 + ["目标", "总结"]
    scenes = [_generation_scene(index, role) for index, role in enumerate(roles, 1)]

    selected = _ai_visual_scene_ids(scenes, "hybrid_director")

    assert len(selected) == 8
    assert not any(scenes[index - 1].pedagogical_role in {"目标", "总结"} for index in range(11, 13) if scenes[index - 1].id in selected)
    shots = _shot_plan(scenes[0], enhanced=True, visual_mode="hybrid_director", ai_asset_id="image-1")
    assert [shot.source_type for shot in shots] == ["ppt", "ai_image", "ppt"]
    assert shots[-1].end_offset_seconds == 10


def test_ai_visual_first_uses_generated_image_for_whole_scene():
    scene = _generation_scene(1, "情境")
    shots = _shot_plan(scene, enhanced=True, visual_mode="ai_visual_first", ai_asset_id="image-1")
    assert len(shots) == 1
    assert shots[0].source_type == "ai_image"
    assert shots[0].start_offset_seconds == 0
    assert shots[0].end_offset_seconds == 10


def test_subtitles_are_split_into_short_timed_cues():
    scene = _generation_scene(1, "概念讲解")
    scene.subtitle_text = "理解这一部分的关键，不是孤立记忆词语，而是看清概念、适用条件和彼此关系。请结合画面依次关注测物重与浸没拉力。"
    srt, vtt = _subtitle_documents([scene])
    cues = [block for block in srt.strip().split("\n\n") if block]
    assert len(cues) >= 3
    assert all(max(map(len, block.splitlines()[2:])) <= 16 for block in cues)
    assert "00:00:10.000" in vtt


def test_image_timeout_is_retryable_and_keeps_a_readable_error():
    error = TimeoutError()
    assert _retryable_image_error(error)
    assert _media_error_detail(error) == "TimeoutError"
    assert _retryable_image_error(RuntimeError("429 Too Many Requests"))
    assert _retryable_image_error(RuntimeError("502 Bad Gateway"))
    assert not _retryable_image_error(ValueError("invalid image payload"))


def test_video_visual_prompt_uses_semantics_without_ppt_layout_instructions():
    prompt = _visual_prompt(
        {
            "slide_id": "S01",
            "learning_purpose": "解释物体上下表面的压力差产生浮力",
            "visual_track": {"composition": "左右分栏，三个带标题的步骤框"},
            "production_notes": ["卡片式布局并展示公式"],
            "audio_track": {"narration_text": "观察浸没方块上下表面的水压方向。"},
        },
        {"slides": [{
            "id": "S01", "title": "浮力产生的原因", "purpose": "观察压力差",
            "visual_suggestion": "左文右图并配大标题",
        }]},
    )

    assert "解释物体上下表面的压力差产生浮力" in prompt
    assert "三个带标题的步骤框" not in prompt
    assert "左文右图并配大标题" not in prompt
    assert "不得出现任何可读字符" in prompt


def test_recompose_refreshes_legacy_visual_prompt_from_latest_sources():
    bp = make_blueprint(sample_course())
    lesson, ppt = make_lesson_plan(bp), make_ppt(bp)
    script = make_video_script(bp, lesson, ppt)
    scene = _generation_scene(1, "导入")
    scene.script_scene_id = script.scenes[0].id
    scene.visual_prompt = "旧版 PPT 左右分栏布局"

    _hydrate_scene_director_data(
        [scene],
        SimpleNamespace(content_json=script.model_dump()),
        SimpleNamespace(content_json=ppt.model_dump()),
        refresh_visual_prompts=True,
    )

    assert scene.visual_prompt != "旧版 PPT 左右分栏布局"
    assert "不得生成 PPT" in scene.visual_prompt


def test_video_generation_v1_is_upgraded_in_memory():
    legacy = {
        "schema_version": "1.0", "mode": "hybrid", "production_settings": {},
        "source_versions": {}, "scenes": [_generation_scene(1, "导入").model_dump()],
        "outputs": {"duration_seconds": 10},
    }
    content = VideoGenerationContent.model_validate(legacy)
    assert content.schema_version == "2.0"
    assert content.production_settings.visual_mode == "ai_visual_first"
    assert content.generation_warnings


@pytest.mark.asyncio
async def test_local_tts_fallback_produces_non_silent_pcm(tmp_path):
    if not _local_tts_engine():
        pytest.skip("当前系统没有本地 TTS")
    output = tmp_path / "narration.wav"
    provider, model = await _generate_local_speech(
        "同学们好，今天我们学习浮力。", output, 4, "natural", 4,
    )
    assert provider and model == "local-tts"
    assert output.stat().st_size > 1024
    assert _pcm_rms(output) >= 80


@pytest.mark.asyncio
async def test_completed_video_run_is_idempotent(monkeypatch, client, auth_headers):
    from app.core.database import SessionLocal
    from app.models.entities import CourseProject, CourseTask, GenerationRun
    from app.services.video_generation_service import execute_video_generation_run

    created = await client.post("/api/v1/courses", headers=auth_headers, json={
        "title": "幂等测试", "subject": "物理", "grade_level": "八年级",
        "audience": "学生", "duration_minutes": 1, "scenario": "课堂",
        "course_task": "验证重复执行不会覆盖成功状态",
    })
    assert created.status_code == 201
    async with SessionLocal() as db:
        course = await db.get(CourseProject, created.json()["id"])
        task = CourseTask(
            course_id=course.id, task_type="video_generation", agent_type="video",
            display_order=6, status="review", progress=100, dependency_types_json=[],
        )
        db.add(task)
        await db.flush()
        run = GenerationRun(
            course_id=course.id, course_task_id=task.id, thread_id="completed-run",
            run_type="task", trigger_type="recompose", status="completed", progress=100,
        )
        db.add(run)
        await db.commit()
        run_id = run.id

    monkeypatch.setattr(
        "app.services.video_generation_service._validate_sources",
        lambda *_args, **_kwargs: (_ for _ in ()).throw(AssertionError("completed run must not execute")),
    )
    await execute_video_generation_run(run_id)

    async with SessionLocal() as db:
        persisted = await db.get(GenerationRun, run_id)
        assert persisted.status == "completed"
